import os
import re
import json
import time
from docx import Document
from src.agents.llm import call_groq_llm

TEMPLATE_DIR = "dichvucong_xay_dung_crawled_2026-07-17/mau_don_to_khai"
OUTPUT_JSON_PATH = "config/standardized_form_fields.json"

# Key forms to prioritize for standardization
KEY_FORMS = [
    "Don2.docx",
    "DDNCapGPXD (1).docx",
    "DDNCapGPXD.docx",
    "Mẫu số 29.docx",
    "Mau so 18_73_2026_QD-UBND.docx",
    "Mẫu số_ 18.docx",
    "Mẫu số 28. đơn biến động đất đai.docx",
    "mẫu số 28.docx"
]

def extract_title(doc, filename):
    """Extract document title from the first few paragraphs of the DOCX."""
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        if not text:
            continue
        if "CỘNG HÒA" in text or "Độc lập" in text or "---" in text:
            continue
        if text.startswith("(") and text.endswith(")"):
            continue
        if any(kw in text.upper() for kw in ["ĐƠN", "GIẤY", "MẪU", "BIÊN BẢN", "TỜ KHAI", "PHIẾU", "BÁO CÁO", "VĂN BẢN"]):
            return text
        upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if upper_ratio > 0.5:
            return text
    return filename.replace(".docx", "")

def extract_label(text):
    """Extract clean label before the blank spaces."""
    # Remove numbering like 1., 2.1., etc.
    cleaned = re.sub(r'^[\d\.]+\s*', '', text.strip())
    # Remove letter numbering like a), b), c), d)
    cleaned = re.sub(r'^[a-z]\)\s*', '', cleaned)
    # Remove leading dashes/bullet points
    cleaned = re.sub(r'^[-\*\+]\s*', '', cleaned)
    
    # Split by blank pattern and get prefix
    parts = re.split(r'[\.]{3,}|[…]{2,}|[_]{3,}', cleaned)
    label = parts[0].strip().rstrip(':').strip()
    
    # Remove footnotes/brackets like (2), (3), (nếu có)
    label = re.sub(r'\(\d+\)', '', label)
    label = re.sub(r'\d+\)', '', label)
    
    label = label.strip()
    if len(label) < 2 or label in ("---", "***", "===", "Kính gửi"):
        return None
        
    # Clean up trailing special chars
    label = label.rstrip('.').rstrip(';').rstrip(':').strip()
    
    # Filter out signature blocks and date placeholders
    if any(w in label.lower() for w in ["ngày … tháng", "ngày... tháng", "ngày tháng năm", "ngày ... tháng"]):
        return None
    if any(w in label.lower() for w in ["người viết đơn", "người nộp đơn", "cam đoan", "hướng dẫn", "ghi thông tin"]):
        return None
        
    # Filter out extremely long instruction text
    if len(label) > 60:
        return None
        
    return label

def extract_blank_fields(doc_path):
    """Scan paragraphs and tables for blank spaces, supporting pure blank paragraphs with previous label heuristic."""
    doc = Document(doc_path)
    blank_fields = []
    blank_pattern = re.compile(r'[\.]{3,}|[…]{2,}|[_]{3,}')
    
    # 1. Paragraphs
    prev_text = ""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        is_pure_blank = len(text) > 3 and all(c in "._… " for c in text)
        if is_pure_blank and prev_text and not blank_pattern.search(prev_text):
            label = extract_label(prev_text)
            if label and not any(f["label"] == label for f in blank_fields):
                blank_fields.append({
                    "label": label,
                    "original_text": prev_text + " " + text
                })
        elif blank_pattern.search(text):
            label = extract_label(text)
            if label and not any(f["label"] == label for f in blank_fields):
                blank_fields.append({
                    "label": label,
                    "original_text": text
                })
        prev_text = text
                
    # 2. Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue
                    if blank_pattern.search(text):
                        label = extract_label(text)
                        if label:
                            # Avoid duplicate labels from merged cells
                            if not any(f["label"] == label for f in blank_fields):
                                blank_fields.append({
                                    "label": label,
                                    "original_text": text
                                })
    return blank_fields

def generate_fields_json(blank_fields, title):
    """Use Gemini to map labels to keys and clean Vietnamese questions."""
    if not blank_fields:
        return []
        
    labels = [f["label"] for f in blank_fields]
    fields_text = "\n".join([f"- {lbl}" for lbl in labels])
    
    system_prompt = f"""Bạn là chuyên gia số hóa biểu mẫu hành chính công.
Tiêu đề biểu mẫu: {title}

Nhiệm vụ của bạn là tạo các trường câu hỏi để điền vào tất cả các khoảng trống trong đơn.
Hãy:
1. Giữ lại tất cả các trường thông tin cần điền bao gồm: Thông tin cá nhân, Số điện thoại, CCCD, Địa chỉ, các thông tin về Giấy chứng nhận đã cấp (Số vào sổ, Ngày cấp, Số phát hành...), Nội dung biến động, Thông tin thửa đất, các nội dung đề nghị khác.
2. Chỉ loại bỏ các khoảng trống dành cho: Chữ ký/Ký tên ở cuối đơn (ví dụ: Người viết đơn ký tên...), Phần xác nhận/phê duyệt của Cơ quan nhà nước ở cuối đơn.
3. Thiết lập mã khóa (key) tiếng Anh ngắn gọn và câu hỏi thu thập thông tin tự nhiên bằng tiếng Việt cho từng trường.

Trả về kết quả dưới dạng JSON object:
{{
  "fields": [
    {{"field_label": "Nhãn trường gốc", "question": "Câu hỏi tự nhiên cho người dân", "key": "khoa_tieng_anh"}}
  ]
}}

Không giải thích thêm, chỉ trả về JSON.

Danh sách các trường cần xử lý:
{fields_text}"""

    api_messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Hãy sinh JSON các câu hỏi tương ứng."}
    ]
    
    # Try calling LLM (Gemini 3.5 Flash)
    response = call_groq_llm(api_messages, temperature=0.1, max_tokens=2048)
    
    try:
        data = json.loads(response)
        if isinstance(data, dict) and "fields" in data:
            return data["fields"]
    except Exception as e:
        print(f"Error parsing LLM response for {title}: {e}")
        
    # Return placeholder fields if failed
    result = []
    for lbl in labels[:10]:
        key = re.sub(r'\W+', '_', lbl).lower().strip('_')
        result.append({
            "field_label": lbl,
            "question": f"Vui lòng nhập thông tin cho: {lbl}?",
            "key": key
        })
    return result

def main():
    print("=== STARTING FORM STANDARDIZATION ===")
    
    # Load existing fields database if it exists
    db = {}
    if os.path.exists(OUTPUT_JSON_PATH):
        try:
            with open(OUTPUT_JSON_PATH, 'r', encoding='utf-8') as f:
                db = json.load(f)
        except Exception:
            db = {}
            
    # Scan files in templates directory
    files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx") and not f.startswith("~$")]
    
    # Prioritize Key Forms
    files_to_scan = [f for f in KEY_FORMS if f in files]
    
    print(f"Total files in folder: {len(files)}")
    print(f"Prioritizing scan for key files: {files_to_scan}")
    
    for filename in files_to_scan:
        file_path = os.path.join(TEMPLATE_DIR, filename)
        print(f"\nProcessing: {filename}...")
        
        try:
            doc = Document(file_path)
            title = extract_title(doc, filename)
            print(f"Extracted Title: {title}")
            
            blank_fields = extract_blank_fields(file_path)
            print(f"Found {len(blank_fields)} blank field placeholders.")
            
            fields_json = generate_fields_json(blank_fields, title)
            print(f"Generated {len(fields_json)} standardized fields.")
            
            db[filename] = {
                "title": title,
                "fields": fields_json
            }
            
            # Save immediately to prevent data loss
            os.makedirs(os.path.dirname(OUTPUT_JSON_PATH), exist_ok=True)
            with open(OUTPUT_JSON_PATH, 'w', encoding='utf-8') as f:
                json.dump(db, f, ensure_ascii=False, indent=2)
                
            # Prevent hitting rate limits
            time.sleep(2)
            
        except Exception as e:
            print(f"Failed standardizing {filename}: {e}")
            
    print(f"\n=== SUCCESS! Standardized JSON database saved to {OUTPUT_JSON_PATH} ===")

if __name__ == "__main__":
    main()
