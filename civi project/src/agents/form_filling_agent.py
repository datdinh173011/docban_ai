import os
import re
import json
from docx import Document
from src.ingestion.metadata_loader import MetadataLoader
from src.agents.llm import call_groq_llm

_CONFIG_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'config', 'standardized_form_fields.json')


class FormFillingAgent:
    def __init__(self):
        self.metadata_loader = MetadataLoader()
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Load standardized forms database if it exists
        self.standardized_db = {}
        if os.path.exists(_CONFIG_PATH):
            try:
                with open(_CONFIG_PATH, 'r', encoding='utf-8') as f:
                    self.standardized_db = json.load(f)
            except Exception as e:
                print(f"[FormFillingAgent Init Error] Failed loading standardized database: {e}")

    def _extract_title(self, doc):
        """Extract the document title from the first few paragraphs of the DOCX."""
        for para in doc.paragraphs[:10]:
            text = para.text.strip()
            # Skip empty, header republic line, independence line, dashes
            if not text:
                continue
            if "CỘNG HÒA" in text or "Độc lập" in text or "---" in text:
                continue
            if text.startswith("(") and text.endswith(")"):
                continue
            if any(kw in text.upper() for kw in ["ĐƠN", "GIẤY", "MẪU", "BIÊN BẢN", "TỜ KHAI", "PHIẾU", "BÁO CÁO", "VĂN BẢN"]):
                return text
            upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
            if upper_ratio > 0.5:
                return text
        return None

    def _extract_blank_fields(self, doc):
        """
        Scan all paragraphs and tables in the DOCX and extract fields that have blank placeholders.
        Returns a list of dicts: [{"original_text": str, "label": str}]
        """
        blank_fields = []
        # Patterns that indicate a blank field: ..., ……, ___,  : …
        blank_pattern = re.compile(r'[\.]{3,}|[…]{2,}|[_]{3,}')
        
        # 1. Scan paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            if blank_pattern.search(text):
                label = self._extract_label(text)
                if label:
                    blank_fields.append({
                        "label": label,
                        "original_text": text
                    })
                    
        # 2. Scan tables (crucial for templates where forms are built inside grids)
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        if not text:
                            continue
                        if blank_pattern.search(text):
                            label = self._extract_label(text)
                            if label:
                                # Prevent duplicates from merged cells
                                if not any(f["label"] == label for f in blank_fields):
                                    blank_fields.append({
                                        "label": label,
                                        "original_text": text
                                    })
        return blank_fields

    def _extract_label(self, text):
        """Extract a meaningful label from a paragraph containing blanks."""
        # Remove numbering at start (e.g., "1.", "- ", "* ")
        cleaned = re.sub(r'^[\d]+[\.\)]\s*', '', text.strip())
        cleaned = re.sub(r'^[-\*]\s*', '', cleaned)
        
        # Split by blank pattern and get prefix
        parts = re.split(r'[\.]{3,}|[…]{2,}|[_]{3,}', cleaned)
        label = parts[0].strip().rstrip(':').strip()
        
        # Skip very short or structural lines
        if len(label) < 3 or label in ("---", "***", "===", "Kính gửi"):
            return None
        return label

    def _generate_questions_from_fields(self, blank_fields, template_title):
        """Use LLM to generate clean Vietnamese questions for dynamically scanned fields."""
        fields_text = "\n".join([f"- {f['label']}" for f in blank_fields])
        
        system_prompt = f"""Bạn là trợ lý giúp tạo câu hỏi để thu thập thông tin điền vào mẫu đơn hành chính.

Tiêu đề mẫu đơn: {template_title}

Dưới đây là danh sách các trường trống trong đơn. Hãy:
1. Chỉ giữ lại các trường mà NGƯỜI DÂN cần tự điền (tên, CCCD, địa chỉ, số điện thoại, thông tin thửa đất, số tầng, diện tích...).
2. Bỏ các trường do CƠ QUAN NHÀ NƯỚC điền (kính gửi, quyết định số, tên tổ chức lập thiết kế, mã chứng chỉ hành nghề, cốt xây dựng, khoảng lùi...) hoặc các trường quá chuyên môn mà người dân thường không biết.
3. Với mỗi trường giữ lại, tạo 1 câu hỏi tự nhiên bằng tiếng Việt.

Trả về đối tượng JSON có cấu trúc sau:
{{
  "fields": [
    {{"field_label": "nhãn gốc từ đơn", "question": "câu hỏi tiếng Việt tự nhiên", "key": "mã ngắn gọn dùng làm khóa"}}
  ]
}}

Chỉ trả về JSON, không giải thích thêm.

Danh sách trường trống:
{fields_text}"""

        api_messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Hãy tạo câu hỏi cho các trường trống trong đơn '{template_title}'."}
        ]
        
        response = call_groq_llm(api_messages, temperature=0.1, max_tokens=2048)
        
        try:
            data = json.loads(response)
            if isinstance(data, dict) and "fields" in data:
                return data["fields"]
        except (json.JSONDecodeError, TypeError):
            pass
        
        return []

    def _save_to_standardized_db(self, template_name, title, fields):
        """Save a newly scanned form configuration to standardized database."""
        self.standardized_db[template_name] = {
            "title": title,
            "fields": fields
        }
        try:
            os.makedirs(os.path.dirname(_CONFIG_PATH), exist_ok=True)
            with open(_CONFIG_PATH, 'w', encoding='utf-8') as f:
                json.dump(self.standardized_db, f, ensure_ascii=False, indent=2)
            print(f"[FormFillingAgent] Saved standardized config for {template_name} to disk.")
        except Exception as e:
            print(f"[FormFillingAgent] Failed writing standardized database: {e}")

    def run(self, state):
        messages = state.get("messages", [])
        last_message = messages[-1].content if messages else ""
        
        selected_proc = state.get("selected_procedure", {})
        if not selected_proc:
            return {"response": "Không thể điền đơn vì chưa xác định được thủ tục."}
            
        # Get form templates for this procedure
        forms = self.metadata_loader.get_forms_for_procedure(selected_proc["code"])
        if not forms:
            return {"response": "Không tìm thấy biểu mẫu DOCX nào tương ứng với thủ tục này để điền tự động."}
            
        form_info = forms[0]
        template_name = form_info.get("FileName")
        template_path = os.path.join(
            "dichvucong_xay_dung_crawled_2026-07-17",
            "mau_don_to_khai",
            template_name
        )
        
        if not os.path.exists(template_path):
            return {"response": f"Không tìm thấy file biểu mẫu: {template_name} trên máy."}
        
        # Load the DOCX template
        doc = Document(template_path)
        
        # Determine title and fields: Check database first, else scan dynamically
        template_title = None
        form_questions = state.get("form_questions", [])
        
        if not form_questions:
            # Check standardized database cache
            if template_name in self.standardized_db:
                print(f"[FormFillingAgent] Loading standardized configuration for {template_name} from database...")
                cached_form = self.standardized_db[template_name]
                template_title = cached_form.get("title")
                form_questions = cached_form.get("fields", [])
            else:
                # Dynamic scan as fallback
                print(f"[FormFillingAgent] Form {template_name} not found in database. Performing dynamic scanning (paragraphs & tables)...")
                template_title = self._extract_title(doc) or template_name
                blank_fields = self._extract_blank_fields(doc)
                form_questions = self._generate_questions_from_fields(blank_fields, template_title)
                
                # Save to database to make it persistent and fast next time
                if form_questions:
                    self._save_to_standardized_db(template_name, template_title, form_questions)
        
        if not template_title:
            template_title = self.standardized_db.get(template_name, {}).get("title") or self._extract_title(doc) or template_name
            
        if not form_questions:
            return {"response": f"Không tìm thấy trường trống nào trong mẫu đơn **{template_title}** để điền."}
        
        user_info = state.get("user_info", {})
        
        # Find current missing field
        current_missing = None
        for field in form_questions:
            key = field.get("key", "")
            if key and key not in user_info:
                current_missing = field
                break
        
        # Check cancel/skip commands
        if last_message.strip().lower() in ("hủy điền đơn", "hủy"):
            return {
                "user_info": {},
                "form_questions": [],
                "filled_form_path": "",
                "fill_form_approved": False,
                "response": "Đã hủy điền đơn mẫu biểu. Tôi có thể hỗ trợ gì khác cho bạn?",
                "suggestions": ["Tôi muốn xây dựng nhà ở", "Tôi muốn đăng ký đất đai", "Tôi muốn chuyển nhượng nhà đất", "Tôi muốn hỏi thông tin thủ tục"]
            }

        # Save user answer
        if current_missing and len(messages) > 1 and not last_message.startswith("/"):
            val = last_message.strip()
            if val.lower() in ("bỏ qua trường này", "bỏ qua"):
                val = "..."
            user_info[current_missing["key"]] = val
            # Find next missing
            current_missing = None
            for field in form_questions:
                key = field.get("key", "")
                if key and key not in user_info:
                    current_missing = field
                    break
        
        # If still have missing fields, ask the next question
        if current_missing:
            return {
                "user_info": user_info,
                "form_questions": form_questions,
                "response": f"Để hoàn thành **{template_title}**, vui lòng cho tôi biết:\n👉 **{current_missing['question']}**",
                "suggestions": ["Bỏ qua trường này", "Hủy điền đơn"]
            }
        
        # All fields collected — perform DOCX filling
        filled_docx_name = f"filled_{selected_proc['code']}_{user_info.get('fullname', user_info.get(list(user_info.keys())[0] if user_info else 'khach', 'khach'))}.docx"
        filled_docx_path = os.path.join(self.output_dir, filled_docx_name)
        
        self._fill_docx_dynamic(doc, filled_docx_path, user_info, form_questions)
        
        return {
            "user_info": user_info,
            "form_questions": form_questions,
            "filled_form_path": filled_docx_path,
            "fill_form_approved": False,  # Reset approval flag after completion
            "response": f"🎉 Tuyệt vời! Tôi đã điền xong thông tin vào mẫu đơn **{template_title}**.\n\n📂 **File kết quả đã được lưu tại:**\n`{os.path.abspath(filled_docx_path)}`\n\nBạn có thể tải file này về máy để chuẩn bị nộp hồ sơ.",
            "suggestions": ["Tôi muốn xây dựng nhà ở", "Tôi muốn đăng ký đất đai", "Tôi muốn chuyển nhượng nhà đất", "Tôi muốn hỏi thông tin thủ tục"]
        }

    def _fill_docx_dynamic(self, doc, output_path, user_info, form_questions):
        """Fill the DOCX by replacing blank patterns in fields that match the collected user info."""
        label_to_value = {}
        for q in form_questions:
            key = q.get("key", "")
            label = q.get("field_label", "")
            if key in user_info and label:
                label_to_value[label] = user_info[key]
        
        blank_pattern = re.compile(r'[\.]{3,}|[…]{2,}|[_]{3,}')
        
        # 1. Fill paragraphs
        for para in doc.paragraphs:
            text = para.text
            if not blank_pattern.search(text):
                continue
            
            for label, value in label_to_value.items():
                if label.lower() in text.lower() or any(word in text.lower() for word in label.lower().split() if len(word) > 3):
                    new_text = blank_pattern.sub(f" {value} ", text, count=1)
                    new_text = re.sub(r'\s{2,}', ' ', new_text).strip()
                    para.text = new_text
                    break
        
        # 2. Fill tables (crucial for grid-aligned text entries)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if not blank_pattern.search(text):
                            continue
                        for label, value in label_to_value.items():
                            if label.lower() in text.lower():
                                new_text = blank_pattern.sub(f" {value} ", text, count=1)
                                new_text = re.sub(r'\s{2,}', ' ', new_text).strip()
                                para.text = new_text
                                break
        
        doc.save(output_path)
