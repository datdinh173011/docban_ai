import os
import re
from docx import Document

def extract_label(text):
    # Remove numbering like 1., 2.1., etc.
    cleaned = re.sub(r'^[\d\.]+\s*', '', text.strip())
    # Remove letter numbering like a), b), c), d)
    cleaned = re.sub(r'^[a-z]\)\s*', '', cleaned)
    # Remove leading dashes/bullet points
    cleaned = re.sub(r'^[-\*\+]\s*', '', cleaned)
    
    # Split by blank pattern and get prefix
    parts = re.split(r'[\.]{3,}|[…]{2,}|[_]{3,}', cleaned)
    label = parts[0].strip().rstrip(':').strip()
    
    # Remove footnotes/brackets like (2), (3), (nếu có)
    label = re.sub(r'\(\d+\)', '', label)
    label = re.sub(r'\d+\)', '', label)
    
    label = label.strip()
    if len(label) < 2 or label in ("---", "***", "===", "Kính gửi"):
        return None
        
    # Clean up trailing special chars
    label = label.rstrip('.').rstrip(';').rstrip(':').strip()
    
    # Filter out signature blocks and date placeholders
    if any(w in label.lower() for w in ["ngày … tháng", "ngày... tháng", "ngày tháng năm", "ngày ... tháng"]):
        return None
    if any(w in label.lower() for w in ["người viết đơn", "người nộp đơn", "cam đoan", "hướng dẫn", "ghi thông tin"]):
        return None
        
    # Filter out extremely long instruction text
    if len(label) > 60:
        return None
        
    return label

def extract_blank_fields(doc_path):
    doc = Document(doc_path)
    blank_fields = []
    blank_pattern = re.compile(r'[\.]{3,}|[…]{2,}|[_]{3,}')
    
    # 1. Paragraphs
    prev_text = ""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        is_pure_blank = len(text) > 3 and all(c in "._… " for c in text)
        if is_pure_blank and prev_text and not blank_pattern.search(prev_text):
            label = extract_label(prev_text)
            if label and not any(f["label"] == label for f in blank_fields):
                blank_fields.append({
                    "label": label,
                    "original_text": prev_text + " " + text
                })
        elif blank_pattern.search(text):
            label = extract_label(text)
            if label and not any(f["label"] == label for f in blank_fields):
                blank_fields.append({
                    "label": label,
                    "original_text": text
                })
        prev_text = text
                
    # 2. Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue
                    if blank_pattern.search(text):
                        label = extract_label(text)
                        if label:
                            if not any(f["label"] == label for f in blank_fields):
                                blank_fields.append({
                                    "label": label,
                                    "original_text": text
                                })
    return blank_fields

fields = extract_blank_fields('dichvucong_xay_dung_crawled_2026-07-17/mau_don_to_khai/Mẫu số 29.docx')
for f in fields:
    print(f"- {f['label']}")
